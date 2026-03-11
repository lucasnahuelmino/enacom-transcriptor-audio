from __future__ import annotations

import datetime as dt
from collections import Counter, defaultdict
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook, load_workbook

from app.core.paths import LOGO_PATH, ensure_dirs
from app.core.utils import sec_to_hhmmss, fmt_hhmmss


# =========================
# Excel helpers
# =========================

def ensure_excel_file(xlsx_path: str, sheets: dict[str, tuple[str, ...] | None]) -> None:
    ensure_dirs()
    p = Path(xlsx_path)
    p.parent.mkdir(parents=True, exist_ok=True)

    if p.exists():
        wb = load_workbook(str(p))
    else:
        wb = Workbook()

    for sheet_name, headers in sheets.items():
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
        else:
            if wb.sheetnames == ["Sheet"] and wb["Sheet"]["A1"].value is None:
                ws = wb["Sheet"]
                ws.title = sheet_name
            else:
                ws = wb.create_sheet(sheet_name)

        if headers is not None and ws["A1"].value is None:
            ws.append(list(headers))

    wb.save(str(p))


def append_to_excel(
    xlsx_path: str,
    row: list,
    sheet_name: str = "Transcripción",
    headers: tuple[str, ...] = ("Inicio", "Fin", "Texto"),
) -> None:
    ensure_excel_file(xlsx_path, {sheet_name: headers})

    p = Path(xlsx_path)
    wb = load_workbook(str(p))
    ws = wb[sheet_name]
    ws.append(row)
    wb.save(str(p))


def write_infracciones_excel(
    xlsx_path: str,
    infracciones: list[dict] | None,
    sheet_name: str = "Infracciones",
    resumen_sheet: str = "Resumen_infracciones",
) -> None:
    ensure_excel_file(xlsx_path, {sheet_name: ("Archivo", "Término", "Inicio", "Fin", "Texto")})

    p = Path(xlsx_path)
    wb = load_workbook(str(p))

    if sheet_name in wb.sheetnames:
        wb.remove(wb[sheet_name])
    ws = wb.create_sheet(sheet_name)
    ws.append(["Archivo", "Término", "Inicio", "Fin", "Texto"])

    infracciones = infracciones or []
    for inf in infracciones:
        ws.append([
            inf.get("archivo", ""),
            inf.get("termino", ""),
            inf.get("inicio", ""),
            inf.get("fin", ""),
            inf.get("texto", ""),
        ])

    if resumen_sheet in wb.sheetnames:
        wb.remove(wb[resumen_sheet])
    ws2 = wb.create_sheet(resumen_sheet)
    ws2.append(["Archivo", "Término", "Ocurrencias"])

    counter = defaultdict(Counter)
    for inf in infracciones:
        counter[inf.get("archivo", "")][inf.get("termino", "")] += 1

    for archivo, c in counter.items():
        for termino, n in c.most_common():
            ws2.append([archivo, termino, n])

    wb.save(str(p))


# =========================
# DOCX helpers
# =========================

def _template_path() -> Path:
    return LOGO_PATH.parent / "plantilla_enacom_limpia.docx"


def _load_doc() -> Document:
    tp = _template_path()
    if tp.exists():
        try:
            return Document(str(tp))
        except Exception:
            pass
    return Document()


def _add_heading_safe(doc: Document, text: str, level: int) -> None:
    style_name = f"Heading {level}"
    try:
        doc.add_paragraph(text, style=style_name)
    except Exception:
        p = doc.add_paragraph(text)
        if p.runs:
            p.runs[0].bold = True


def _kv_table(doc: Document, rows: Iterable[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"

    for k, v in rows:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)


def generar_informe_word(
    titulo: str,
    docx_out_path: str,
    combinado: bool,
    meta: dict,
    whisper_result: dict | None = None,
    infracciones: list[dict] | None = None,
    files_info: list[dict] | None = None,
) -> str:
    ensure_dirs()

    doc = _load_doc()

    _add_heading_safe(doc, "Transcriptor de audios a texto", level=1)
    doc.add_paragraph("Informe combinado (lote)" if combinado else "Informe individual")
    doc.add_paragraph("")

    _add_heading_safe(doc, "Datos del procesamiento", level=2)

    generado = meta.get("generado") or dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # FIX 1: incluir Referencia/Nombre en la tabla de metadatos del informe Word.
    referencia = meta.get("referencia", "") or "—"

    _kv_table(
        doc,
        [
            ("Referencia / Nombre",   referencia),
            ("Archivo / Lote",        titulo),
            ("Generado",              generado),
            ("Modelo Whisper",        meta.get("model_size", "")),
            ("Idioma",                meta.get("lang", "auto") or "auto"),
            ("Duración segmento (s)", meta.get("segment_duration", "")),
            ("Cantidad de archivos",  meta.get("total_files", "")),
            # FIX 2: total_duration_hhmmss ya viene calculado con dur_final desde
            # processing.py — ahora refleja la duración real y no 00:00:00.
            ("Duración total",        meta.get("total_duration_hhmmss", "")),
        ],
    )

    doc.add_paragraph("")

    def _render_whisper_block(doc: Document, wr: dict) -> None:
        segments = wr.get("segments", [])
        full_text = (wr.get("text") or "").strip()

        _add_heading_safe(doc, "Transcripción segmentada", level=3)

        if not segments:
            doc.add_paragraph("(No hay segmentos disponibles.)")
        else:
            for seg in segments:
                # FIX 3: preferir la clave "line" que viene pre-formateada desde
                # merge_text_into_fixed_windows (formato "[00:MM:SS - 00:MM:SS] texto").
                # Esto evita que los timestamps aparezcan como floats crudos ("420.0")
                # cuando se pasaban los segments_acc raw de Whisper.
                if "line" in seg and seg["line"]:
                    doc.add_paragraph(str(seg["line"]))
                else:
                    # Fallback: segmentos raw de Whisper (float → HH:MM:SS con ceros).
                    ini = fmt_hhmmss(float(seg.get("start", 0) or 0))
                    fin = fmt_hhmmss(float(seg.get("end",   0) or 0))
                    txt = (seg.get("text", "") or "").strip() or "Silencio"
                    doc.add_paragraph(f"[{ini} - {fin}] {txt}")

        doc.add_paragraph("")
        doc.add_paragraph("—" * 30)
        doc.add_paragraph("")

        _add_heading_safe(doc, "Transcripción completa", level=3)

        if full_text:
            doc.add_paragraph(full_text)
        else:
            doc.add_paragraph("(No hay texto completo disponible.)")

    # Transcripción
    _add_heading_safe(doc, "Transcripción", level=2)

    if combinado:
        files_info = files_info or []

        if not files_info:
            doc.add_paragraph("(No hay archivos para mostrar.)")
        else:
            for i, info in enumerate(files_info):
                if i > 0:
                    doc.add_page_break()

                archivo = info.get("archivo", f"archivo_{i}")
                dur     = info.get("duracion_hhmmss", "")
                wr      = info.get("whisper_result")

                _add_heading_safe(doc, f"{archivo} — Duración: {dur}", level=3)

                if not wr:
                    doc.add_paragraph("(Sin transcripción disponible.)")
                    continue

                _render_whisper_block(doc, wr)

    else:
        if not whisper_result:
            doc.add_paragraph("(No se recibió resultado de Whisper.)")
        else:
            _render_whisper_block(doc, whisper_result)

    # Infracciones
    doc.add_page_break()
    _add_heading_safe(doc, "Infracciones detectadas", level=2)

    infracciones = infracciones or []

    if not infracciones:
        doc.add_paragraph("No se detectaron infracciones según la configuración aplicada.")
    else:
        resumen = defaultdict(Counter)
        for inf in infracciones:
            resumen[inf.get("archivo", "")][inf.get("termino", "")] += 1

        doc.add_paragraph("Resumen de ocurrencias por archivo y término.")

        t = doc.add_table(rows=1, cols=3)
        h = t.rows[0].cells
        h[0].text = "Archivo"
        h[1].text = "Término"
        h[2].text = "Ocurrencias"

        for archivo, c in resumen.items():
            for termino, n in c.most_common():
                r = t.add_row().cells
                r[0].text = str(archivo)
                r[1].text = str(termino)
                r[2].text = str(n)

        doc.add_paragraph("")
        doc.add_paragraph("Detalle:")

        t2 = doc.add_table(rows=1, cols=5)
        h2 = t2.rows[0].cells
        h2[0].text = "Archivo"
        h2[1].text = "Término"
        h2[2].text = "Inicio"
        h2[3].text = "Fin"
        h2[4].text = "Texto"

        for inf in infracciones[:300]:
            r = t2.add_row().cells
            r[0].text = str(inf.get("archivo", ""))
            r[1].text = str(inf.get("termino", ""))
            r[2].text = str(inf.get("inicio", ""))
            r[3].text = str(inf.get("fin", ""))
            r[4].text = str(inf.get("texto", ""))

    out = Path(docx_out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    return str(out)

def write_combined_txt(txt_path: Path, resultados: list[dict], referencia: str, run_id: str):
    """TXT combinado de todos los archivos"""
    lines = []
    lines.append(f"=== TRANSCRIPCIÓN COMBINADA ===")
    lines.append(f"Referencia: {referencia or '—'}")
    lines.append(f"Generado: {run_id}")
    lines.append("")
    
    for r in resultados:
        lines.append(f"\n{'='*60}")
        lines.append(f"ARCHIVO: {r['archivo']}")
        lines.append(f"Duración: {r['duracion_hhmmss']}")
        lines.append(f"{'='*60}\n")
        lines.append(r['texto_completo'])
    
    txt_path.write_text("\n".join(lines), encoding="utf-8")

def write_combined_xlsx(xlsx_path: Path, files_info: list[dict], infracciones: list[dict]):
    """XLSX combinado con todos los archivos"""
    ensure_excel_file(str(xlsx_path), {"Transcripción": ("Archivo", "Inicio", "Fin", "Texto")})
    
    wb = load_workbook(str(xlsx_path))
    ws = wb["Transcripción"]
    
    for info in files_info:
        wr = info.get("whisper_result", {})
        for seg in wr.get("segments", []):
            ws.append([
                info["archivo"],
                seg.get("line", "").split("]")[0].replace("[", "").split("-")[0].strip(),
                seg.get("line", "").split("]")[0].replace("[", "").split("-")[1].strip() if "-" in seg.get("line", "") else "",
                seg.get("text", "")
            ])
    
    write_infracciones_excel(str(xlsx_path), infracciones)
    wb.save(str(xlsx_path))

def write_combined_docx(docx_path: Path, title: str, meta: dict, files_info: list[dict], infracciones: list[dict]):
    """DOCX combinado"""
    generar_informe_word(
        titulo=title,
        docx_out_path=str(docx_path),
        combinado=True,
        meta=meta,
        files_info=files_info,
        infracciones=infracciones
    )

def create_zip(zip_path: Path, output_dir: Path, resultados: list[dict], lote_urls: dict):
    """Crea ZIP con todos los archivos"""
    import zipfile
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for r in resultados:
            for key in ['txt_url', 'xlsx_url', 'docx_url']:
                url = r.get(key)
                if url:
                    filename = url.split('/')[-1]
                    file_path = output_dir / filename
                    if file_path.exists():
                        zf.write(file_path, f"individuales/{filename}")
        
        for key in ['lote_txt_url', 'lote_xlsx_url', 'lote_docx_url']:
            url = lote_urls.get(key)
            if url:
                filename = url.split('/')[-1]
                file_path = output_dir / filename
                if file_path.exists():
                    zf.write(file_path, f"lote/{filename}")